from __future__ import annotations

import argparse
import json
import sys
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

SCRIPT_DIR = Path(__file__).resolve().parent
if str(SCRIPT_DIR) not in sys.path:
    sys.path.insert(0, str(SCRIPT_DIR))

from docx_writer_core import (  # noqa: E402
    FrontMatterField,
    WriterConfig,
    build_style_map,
    generate_document as generate_document_from_template,
    inspect_template as inspect_template_profile,
    load_json,
)
from docx_writer_template_hooks import (  # noqa: E402
    ParagraphAnchor,
    StaticTemplateBlockHook,
    apply_static_template_block_hooks,
    append_paragraphs_from_prototype,
    find_first_matching_paragraph,
)

SKILL_DIR = SCRIPT_DIR.parent
ADAPTER_ID = "jos"
ADAPTER_DESCRIPTION = (
    "Journal of Software academic DOCX template adapter for Chinese single-column manuscript layouts."
)
SPEC_FAMILY = "docx-report-spec-v1"
DEFAULT_TEMPLATE_PATH = SKILL_DIR / "assets/软件学报排版样例2025年版.docx"
DEFAULT_PROFILE_PATH = SKILL_DIR / "references/jos-template-profile.json"


STYLE_CANDIDATES: dict[str, list[str]] = {
    "title_cn": ["Subtitle", "Title", "Normal"],
    "authors_cn": ["作者", "author", "Normal"],
    "affiliations_cn": ["单位", "authorinfo", "Normal"],
    "abstract_cn": ["摘要", "abstract", "Body Text", "Normal"],
    "keywords_cn": ["关键词", "Body Text", "Normal"],
    "classification_cn": ["Title", "分类号", "Normal"],
    "citation_cn": ["Name", "文前文本", "Normal"],
    "citation_en": ["Depart.Correspond.http", "Depart.Correspond", "Normal"],
    "title_en": ["Title", "Title1", "Normal"],
    "authors_en": ["Name", "author", "Normal"],
    "affiliations_en": ["Depart.Correspond.http", "Depart.Correspond", "Affiliation", "Normal"],
    "abstract_en": ["Abstract", "abstract", "Body Text", "Normal"],
    "keywords_en": ["Date", "Key words", "Normal"],
    "heading_1": ["Heading 1", "1级标题", "标题1", "Normal"],
    "heading_2": ["Heading 2", "2级标题", "标题11", "Normal"],
    "heading_3": ["Heading 3", "3级标题", "Normal"],
    "body": ["样式1", "Body Text First Indent", "Body Text", "Normal"],
    "ack_heading": ["致谢", "Heading 1", "Normal"],
    "reference_heading": ["Reference", "参考文献", "Normal"],
    "reference_cn_appendix_heading": ["中文参考文献", "Reference", "Normal"],
    "reference_cn_text": ["Text of 中文参考文献", "中文参考文献", "Normal"],
    "reference_en_text": ["Text of Reference 1", "Text of Reference", "Reference", "Normal"],
    "author_bio_text": ["Text of 中文参考文献", "Normal"],
    "appendix_heading": ["Heading 1", "1级标题", "Normal"],
}

FRONT_MATTER_FIELDS = (
    FrontMatterField(
        spec_key="title_cn",
        style_key="title_cn",
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    ),
    FrontMatterField(
        spec_key="authors_cn",
        style_key="authors_cn",
        mode="join_list",
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    ),
    FrontMatterField(
        spec_key="affiliations_cn",
        style_key="affiliations_cn",
        mode="list",
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    ),
    FrontMatterField(
        spec_key="abstract_cn",
        style_key="abstract_cn",
        label="摘  要:",
    ),
    FrontMatterField(
        spec_key="keywords_cn",
        style_key="keywords_cn",
        mode="keywords",
        label="关键词:",
        separator=";",
    ),
    FrontMatterField(
        spec_key="classification_cn",
        style_key="classification_cn",
        label="中图法分类号:",
    ),
    FrontMatterField(
        spec_key="citation_cn",
        style_key="citation_cn",
    ),
    FrontMatterField(
        spec_key="citation_en",
        style_key="citation_en",
    ),
    FrontMatterField(
        spec_key="title_en",
        style_key="title_en",
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    ),
    FrontMatterField(
        spec_key="authors_en",
        style_key="authors_en",
        mode="join_list",
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    ),
    FrontMatterField(
        spec_key="affiliations_en",
        style_key="affiliations_en",
        mode="list",
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    ),
    FrontMatterField(
        spec_key="abstract_en",
        style_key="abstract_en",
        label="Abstract:",
    ),
    FrontMatterField(
        spec_key="keywords_en",
        style_key="keywords_en",
        mode="keywords",
        label="Key words:",
        separator="; ",
    ),
)

JOS_WRITER_CONFIG = WriterConfig(
    style_candidates=STYLE_CANDIDATES,
    front_matter_fields=FRONT_MATTER_FIELDS,
    body_columns=1,
    acknowledgement_heading="致谢",
    reference_heading="References",
    appendix_default_title="附录",
)

STATIC_TEMPLATE_BLOCK_HOOKS = (
    StaticTemplateBlockHook(
        name="front_masthead",
        placement="prepend",
        mode="before_first_match",
        anchor=ParagraphAnchor(
            style_key="title_cn",
            require_non_empty_text=True,
        ),
    ),
)


def inspect_template(template_path: str | Path, sample_limit: int = 40) -> dict[str, Any]:
    return inspect_template_profile(
        template_path=template_path,
        style_candidates=STYLE_CANDIDATES,
        sample_limit=sample_limit,
    )


def generate_document(
    template_path: str | Path,
    spec: dict[str, Any],
    output_path: str | Path,
) -> dict[str, Any]:
    result = generate_document_from_template(
        template_path=template_path,
        spec=spec,
        output_path=output_path,
        config=JOS_WRITER_CONFIG,
    )
    hook_stats = _apply_dynamic_template_hooks(
        template_path=template_path,
        output_path=output_path,
        spec=spec,
    )
    static_block_counts = apply_static_template_block_hooks(
        template_path=template_path,
        output_path=output_path,
        style_candidates=STYLE_CANDIDATES,
        hooks=STATIC_TEMPLATE_BLOCK_HOOKS,
    )
    hook_stats["static_block_counts"] = static_block_counts
    hook_stats["prepended_front_block_count"] = static_block_counts.get("front_masthead", 0)
    result["template_hook_stats"] = hook_stats
    result["preserved_template_paragraph_count"] = hook_stats["prepended_front_block_count"]
    return result


def _apply_dynamic_template_hooks(
    template_path: str | Path,
    spec: dict[str, Any],
    output_path: str | Path,
) -> dict[str, int]:
    template_doc = Document(str(template_path))
    output_doc = Document(str(output_path))
    template_style_map = build_style_map(template_doc, STYLE_CANDIDATES)
    hook_stats = {
        "inserted_front_matter_block_count": _write_template_specific_front_matter(
            output_doc,
            spec,
            template_doc,
            template_style_map,
        ),
        "appended_back_block_count": _append_static_back_blocks(
            output_doc,
            spec,
            template_doc,
            template_style_map,
        ),
    }

    if any(hook_stats.values()):
        output_doc.save(str(output_path))
    return hook_stats


def _write_template_specific_front_matter(
    doc: Document,
    spec: dict[str, Any],
    template_doc: Document,
    style_map: dict[str, str],
) -> int:
    # Reserved for JOS-only front matter blocks that should be synthesized
    # from spec instead of copied from the template.
    _ = (doc, spec, template_doc, style_map)
    return 0


def _append_static_back_blocks(
    doc: Document,
    spec: dict[str, Any],
    template_doc: Document,
    style_map: dict[str, str],
) -> int:
    sections = _normalize_back_matter_sections(spec)
    if not sections:
        return 0

    count = 0
    for section in sections:
        kind = section["kind"]
        if kind == "references_cn":
            count += _append_cn_reference_appendix(
                doc,
                template_doc,
                style_map,
                title=section["title"],
                items=section["items"],
            )
            continue
        if kind == "author_bios":
            count += _append_author_bio_section(
                doc,
                template_doc,
                style_map,
                title=section["title"],
                items=section["items"],
            )
    return count


def _append_cn_reference_appendix(
    doc: Document,
    template_doc: Document,
    style_map: dict[str, str],
    *,
    title: str,
    items: list[str],
) -> int:
    if not items:
        return 0

    heading_prototype = find_first_matching_paragraph(
        template_doc,
        style_map,
        ParagraphAnchor(text_equals="附中文参考文献"),
    )
    item_prototype = None
    if heading_prototype is not None:
        heading_index = _paragraph_index(template_doc, heading_prototype)
        item_prototype = find_first_matching_paragraph(
            template_doc,
            style_map,
            ParagraphAnchor(
                style_key="reference_cn_text",
                text_contains="[",
                require_non_empty_text=True,
            ),
            start_index=heading_index + 1,
        )

    count = 0
    count += _append_paragraphs_by_prototype_or_style(
        doc,
        prototype_paragraph=heading_prototype,
        style_name=style_map["reference_cn_appendix_heading"],
        texts=[title],
    )
    count += _append_paragraphs_by_prototype_or_style(
        doc,
        prototype_paragraph=item_prototype,
        style_name=style_map["reference_cn_text"],
        texts=items,
    )
    return count


def _append_author_bio_section(
    doc: Document,
    template_doc: Document,
    style_map: dict[str, str],
    *,
    title: str,
    items: list[str],
) -> int:
    if not items:
        return 0

    heading_prototype = find_first_matching_paragraph(
        template_doc,
        style_map,
        ParagraphAnchor(text_equals="作者简介"),
    )
    item_prototype = None
    if heading_prototype is not None:
        heading_index = _paragraph_index(template_doc, heading_prototype)
        item_prototype = find_first_matching_paragraph(
            template_doc,
            style_map,
            ParagraphAnchor(
                style_key="author_bio_text",
                require_non_empty_text=True,
            ),
            start_index=heading_index + 1,
        )

    count = 0
    count += _append_paragraphs_by_prototype_or_style(
        doc,
        prototype_paragraph=heading_prototype,
        style_name=style_map["author_bio_text"],
        texts=[title],
    )
    count += _append_paragraphs_by_prototype_or_style(
        doc,
        prototype_paragraph=item_prototype,
        style_name=style_map["author_bio_text"],
        texts=items,
    )
    return count


def _append_paragraphs_by_prototype_or_style(
    doc: Document,
    *,
    prototype_paragraph,
    style_name: str,
    texts: list[str],
) -> int:
    if prototype_paragraph is not None:
        return append_paragraphs_from_prototype(doc, prototype_paragraph, texts)

    count = 0
    for text in texts:
        content = str(text).strip()
        if not content:
            continue
        doc.add_paragraph(content, style=style_name)
        count += 1
    return count


def _normalize_back_matter_sections(spec: dict[str, Any]) -> list[dict[str, Any]]:
    sections: list[dict[str, Any]] = []
    raw_sections = spec.get("back_matter")
    if isinstance(raw_sections, dict):
        raw_sections = [raw_sections]
    for raw_section in raw_sections or []:
        if not isinstance(raw_section, dict):
            continue
        kind = str(raw_section.get("kind") or raw_section.get("type") or "").strip()
        items = _normalize_text_items(
            raw_section.get("items")
            if raw_section.get("items") is not None
            else raw_section.get("entries")
        )
        if not kind or not items:
            continue
        title = str(raw_section.get("title") or "").strip() or _default_back_matter_title(kind)
        sections.append(
            {
                "kind": kind,
                "title": title,
                "items": items,
            }
        )

    if not any(section["kind"] == "references_cn" for section in sections):
        items = _normalize_text_items(spec.get("references_cn"))
        if items:
            sections.append(
                {
                    "kind": "references_cn",
                    "title": "附中文参考文献",
                    "items": items,
                }
            )

    if not any(section["kind"] == "author_bios" for section in sections):
        items = _normalize_text_items(spec.get("author_bios") or spec.get("author_bios_cn"))
        if items:
            sections.append(
                {
                    "kind": "author_bios",
                    "title": "作者简介",
                    "items": items,
                }
            )

    return sections


def _normalize_text_items(value: Any) -> list[str]:
    if value is None:
        return []
    if isinstance(value, str):
        text = value.strip()
        return [text] if text else []

    items: list[str] = []
    if isinstance(value, list):
        for item in value:
            if isinstance(item, dict):
                text = str(item.get("text") or item.get("content") or "").strip()
            else:
                text = str(item).strip()
            if text:
                items.append(text)
    return items


def _default_back_matter_title(kind: str) -> str:
    if kind == "references_cn":
        return "附中文参考文献"
    if kind == "author_bios":
        return "作者简介"
    return kind


def _paragraph_index(template_doc: Document, target_paragraph) -> int:
    for index, paragraph in enumerate(template_doc.paragraphs):
        if paragraph._p is target_paragraph._p:
            return index
    return -1


def main() -> None:
    parser = argparse.ArgumentParser(description="Inspect and write Journal of Software template DOCX files.")
    subparsers = parser.add_subparsers(dest="command", required=True)

    inspect_parser = subparsers.add_parser("inspect", help="Inspect template styles and structure.")
    inspect_parser.add_argument("--template", required=True, help="Path to the .docx template.")
    inspect_parser.add_argument("--out", help="Optional JSON output path.")

    generate_parser = subparsers.add_parser("generate", help="Generate a .docx from a JSON spec.")
    generate_parser.add_argument("--template", required=True, help="Path to the .docx template.")
    generate_parser.add_argument("--spec", required=True, help="Path to the JSON spec.")
    generate_parser.add_argument("--output", required=True, help="Path to the output .docx.")

    args = parser.parse_args()

    if args.command == "inspect":
        result = inspect_template(args.template)
        text = json.dumps(result, ensure_ascii=False, indent=2)
        if args.out:
            Path(args.out).write_text(text, encoding="utf-8")
        print(text)
        return

    result = generate_document(
        template_path=args.template,
        spec=load_json(args.spec),
        output_path=args.output,
    )
    print(json.dumps(result, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
