from __future__ import annotations

from copy import deepcopy
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml.ns import qn

from docx_writer_core import build_style_map


@dataclass(frozen=True)
class ParagraphAnchor:
    style_key: str | None = None
    text_equals: str | None = None
    text_contains: str | None = None
    require_non_empty_text: bool = False


@dataclass(frozen=True)
class StaticTemplateBlockHook:
    name: str
    placement: str
    mode: str
    anchor: ParagraphAnchor


def apply_static_template_block_hooks(
    template_path: str | Path,
    output_path: str | Path,
    style_candidates: dict[str, list[str]],
    hooks: tuple[StaticTemplateBlockHook, ...],
) -> dict[str, int]:
    if not hooks:
        return {}

    template_doc = Document(str(template_path))
    output_doc = Document(str(output_path))
    style_map = build_style_map(template_doc, style_candidates)

    hook_counts: dict[str, int] = {}
    for hook in hooks:
        paragraphs = _collect_hook_paragraphs(template_doc, style_map, hook)
        if hook.placement == "prepend":
            _prepend_paragraphs(output_doc, paragraphs)
        elif hook.placement == "append":
            _append_paragraphs(output_doc, paragraphs)
        else:
            raise ValueError(f"Unsupported static template hook placement: {hook.placement}")
        hook_counts[hook.name] = len(paragraphs)

    if any(hook_counts.values()):
        output_doc.save(str(output_path))
    return hook_counts


def _collect_hook_paragraphs(
    template_doc: DocxDocument,
    style_map: dict[str, str],
    hook: StaticTemplateBlockHook,
) -> list:
    if hook.mode == "before_first_match":
        return _collect_before_first_match(template_doc, style_map, hook.anchor)
    if hook.mode == "from_first_match":
        return _collect_from_first_match(template_doc, style_map, hook.anchor)
    raise ValueError(f"Unsupported static template hook mode: {hook.mode}")


def _collect_before_first_match(
    template_doc: DocxDocument,
    style_map: dict[str, str],
    anchor: ParagraphAnchor,
) -> list:
    paragraphs = []
    for paragraph in template_doc.paragraphs:
        if _paragraph_matches(paragraph, style_map, anchor):
            break
        paragraphs.append(deepcopy(paragraph._p))
    return paragraphs


def _collect_from_first_match(
    template_doc: DocxDocument,
    style_map: dict[str, str],
    anchor: ParagraphAnchor,
) -> list:
    paragraphs = []
    matched = False
    for paragraph in template_doc.paragraphs:
        if not matched and _paragraph_matches(paragraph, style_map, anchor):
            matched = True
        if matched:
            paragraphs.append(deepcopy(paragraph._p))
    return paragraphs


def _paragraph_matches(paragraph, style_map: dict[str, str], anchor: ParagraphAnchor) -> bool:
    text = paragraph.text.strip()
    if anchor.require_non_empty_text and not text:
        return False
    if anchor.style_key is not None and paragraph.style.name != style_map[anchor.style_key]:
        return False
    if anchor.text_equals is not None and text != anchor.text_equals:
        return False
    if anchor.text_contains is not None and anchor.text_contains not in text:
        return False
    return True


def _prepend_paragraphs(doc: DocxDocument, paragraphs: list) -> None:
    if not paragraphs:
        return
    body = doc._element.body
    for index, paragraph in enumerate(paragraphs):
        body.insert(index, deepcopy(paragraph))


def _append_paragraphs(doc: DocxDocument, paragraphs: list) -> None:
    if not paragraphs:
        return
    body = doc._element.body
    insert_at = len(body)
    if body.sectPr is not None:
        insert_at -= 1
    for offset, paragraph in enumerate(paragraphs):
        body.insert(insert_at + offset, deepcopy(paragraph))


def find_first_matching_paragraph(
    template_doc: DocxDocument,
    style_map: dict[str, str],
    anchor: ParagraphAnchor,
    *,
    start_index: int = 0,
):
    for paragraph in template_doc.paragraphs[start_index:]:
        if _paragraph_matches(paragraph, style_map, anchor):
            return paragraph
    return None


def append_paragraphs_from_prototype(doc: DocxDocument, prototype_paragraph, texts: list[str]) -> int:
    count = 0
    for text in texts:
        content = str(text).strip()
        if not content:
            continue
        paragraph = doc.add_paragraph(style=prototype_paragraph.style.name)
        _copy_paragraph_properties(paragraph, prototype_paragraph)
        run = paragraph.add_run(content)
        _copy_first_run_properties(run, prototype_paragraph)
        count += 1
    return count


def _copy_paragraph_properties(paragraph, prototype_paragraph) -> None:
    source_p_pr = prototype_paragraph._p.find(qn("w:pPr"))
    if source_p_pr is None:
        return
    target_p_pr = paragraph._p.get_or_add_pPr()
    paragraph._p.replace(target_p_pr, deepcopy(source_p_pr))


def _copy_first_run_properties(run, prototype_paragraph) -> None:
    for prototype_run in prototype_paragraph.runs:
        source_r_pr = prototype_run._r.find(qn("w:rPr"))
        if source_r_pr is None:
            continue
        target_r_pr = run._r.get_or_add_rPr()
        run._r.replace(target_r_pr, deepcopy(source_r_pr))
        return
