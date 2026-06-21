from __future__ import annotations

import argparse
import json
import re
import sys
from dataclasses import dataclass
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

SCRIPT_DIR = Path(__file__).resolve().parent
if str(SCRIPT_DIR) not in sys.path:
    sys.path.insert(0, str(SCRIPT_DIR))

from docx_writer_core import (  # noqa: E402
    FrontMatterField,
    WriterConfig,
    generate_document as generate_document_from_template,
    inspect_template as inspect_template_profile,
    load_json,
)


ADAPTER_ID = "generic"
ADAPTER_DESCRIPTION = "Best-effort academic DOCX adapter for user-supplied templates."
SPEC_FAMILY = "docx-report-spec-v1"
DEFAULT_TEMPLATE_PATH = None
DEFAULT_PROFILE_PATH = None

TOP_LEVEL_HEADING_RE = re.compile(r"^\s*\d+\s+\S+")
SECOND_LEVEL_HEADING_RE = re.compile(r"^\s*\d+\.\d+\s+\S+")
THIRD_LEVEL_HEADING_RE = re.compile(r"^\s*\d+\.\d+\.\d+\s+\S+")
REFERENCE_ITEM_RE = re.compile(r"^\s*\[(\d+)\]\s*")


@dataclass(frozen=True)
class TemplateHints:
    preferred_styles: dict[str, str]
    labels: dict[str, str]
    label_joiners: dict[str, str]
    list_joiners: dict[str, str]
    inferred_body_columns: int | None
    reference_heading_text: str | None
    acknowledgement_heading_text: str | None
    reference_numbering_mode: str
    header_sample_text: str | None


BASE_STYLE_CANDIDATES: dict[str, list[str]] = {
    "title_cn": ["Title", "Heading 1", "Subtitle", "Title1", "标题", "题名", "Paper Title", "Normal"],
    "authors_cn": ["人名", "作者", "Author", "Authors", "Subtitle", "Name", "author", "Normal"],
    "affiliations_cn": [
        "地名",
        "单位",
        "Affiliation",
        "authorinfo",
        "作者单位中文",
        "Depart.Correspond",
        "Normal",
    ],
    "abstract_cn": ["摘要", "Abstract", "abstract", "Body Text", "正文文字", "Normal"],
    "keywords_cn": ["关键词", "Key words", "Keywords", "摘要", "Body Text", "Normal"],
    "classification_cn": ["分类号", "摘要", "中图分类号", "Body Text", "Normal"],
    "submission_dates_cn": ["收稿日期", "Information", "正文文字", "Body Text", "Normal"],
    "corresponding_author": ["通信地址", "Correspond", "Information", "正文文字", "Body Text", "Normal"],
    "funding_cn": ["收稿日期", "Information", "正文文字", "Body Text", "Normal"],
    "citation_cn": ["文前文本", "Name", "Body Text", "Normal"],
    "title_en": ["Title", "Heading 1", "Title1", "Subtitle", "English Title", "Paper Title", "Normal"],
    "authors_en": ["人名", "Author", "Authors", "author", "Name", "Subtitle", "Normal"],
    "affiliations_en": ["地名", "Affiliation", "Institution", "Depart.Correspond", "Body Text", "Normal"],
    "abstract_en": ["摘要", "Abstract", "abstract", "Body Text", "Normal"],
    "keywords_en": ["关键词", "Keywords", "Key words", "Date", "Body Text", "Normal"],
    "funding_en": ["Information", "正文文字", "Body Text", "Normal"],
    "citation_en": ["Depart.Correspond.http", "Depart.Correspond", "Affiliation", "Normal"],
    "heading_1": ["Heading 1", "Heading 2", "标题1", "1级标题", "Section", "Normal"],
    "heading_2": ["Heading 2", "Heading 3", "标题11", "2级标题", "Subsection", "Normal"],
    "heading_3": ["Heading 3", "标题111", "3级标题", "Subsubsection", "Normal"],
    "body": ["Body Text", "正文文字", "正文1", "普通正文", "样式1", "Normal"],
    "ack_heading": ["致谢", "Acknowledgements", "Acknowledgments", "Heading 1", "Normal"],
    "reference_heading": ["文献", "参考文献", "参考文献标题", "References", "Reference", "Heading 1", "Normal"],
    "reference_cn_text": [
        "文献文",
        "论文的参考文献",
        "Text of 中文参考文献",
        "中文参考文献",
        "Text of Reference",
        "Reference",
        "Body Text",
        "Normal",
    ],
    "reference_en_text": [
        "文献文",
        "Text of Reference 1",
        "Text of Reference",
        "referenceitem",
        "Reference",
        "Body Text",
        "Normal",
    ],
    "appendix_heading": ["Appendix", "附录", "Heading 1", "Normal"],
}


def inspect_template(template_path: str | Path, sample_limit: int = 40) -> dict[str, Any]:
    doc = Document(str(template_path))
    hints = _collect_template_hints(doc)
    result = inspect_template_profile(
        template_path=template_path,
        style_candidates=_build_style_candidates(hints),
        sample_limit=sample_limit,
    )
    result["template_hints"] = _serialize_hints(hints)
    if hints.inferred_body_columns is not None:
        result["inferred_body_columns"] = hints.inferred_body_columns
    return result


def generate_document(
    template_path: str | Path,
    spec: dict[str, Any],
    output_path: str | Path,
) -> dict[str, Any]:
    template_doc = Document(str(template_path))
    hints = _collect_template_hints(template_doc)
    prepared_spec = _prepare_spec(spec, hints)
    config = _build_writer_config(prepared_spec, hints)
    result = generate_document_from_template(
        template_path=template_path,
        spec=prepared_spec,
        output_path=output_path,
        config=config,
    )
    header_updated = _update_running_headers(template_doc, prepared_spec, output_path)
    result["inferred_body_columns"] = hints.inferred_body_columns
    result["template_hints"] = _serialize_hints(hints)
    result["header_updated"] = header_updated
    return result


def _build_writer_config(spec: dict[str, Any], hints: TemplateHints) -> WriterConfig:
    language = _detect_primary_language(spec)
    style_candidates = _build_style_candidates(hints)
    front_matter_fields = _build_front_matter_fields(hints)
    reference_heading = hints.reference_heading_text or ("References" if language == "en" else "参考文献")
    acknowledgement_heading = hints.acknowledgement_heading_text or (
        "Acknowledgements" if language == "en" else "致谢"
    )
    appendix_default_title = "Appendix" if language == "en" else "附录"
    return WriterConfig(
        style_candidates=style_candidates,
        front_matter_fields=front_matter_fields,
        body_columns=hints.inferred_body_columns,
        acknowledgement_heading=acknowledgement_heading,
        reference_heading=reference_heading,
        appendix_default_title=appendix_default_title,
        reference_numbering_mode=hints.reference_numbering_mode,
    )


def _build_front_matter_fields(hints: TemplateHints) -> tuple[FrontMatterField, ...]:
    def label_and_joiner(spec_key: str, default_label: str, default_joiner: str) -> tuple[str, str]:
        return hints.labels.get(spec_key, default_label), hints.label_joiners.get(spec_key, default_joiner)

    authors_cn_joiner = hints.list_joiners.get("authors_cn", "  ")
    authors_en_joiner = hints.list_joiners.get("authors_en", "  ")
    abstract_cn_label, abstract_cn_joiner = label_and_joiner("abstract_cn", "摘要：", "")
    keywords_cn_label, keywords_cn_joiner = label_and_joiner("keywords_cn", "关键词：", "")
    classification_label, classification_joiner = label_and_joiner("classification_cn", "中图分类号：", "")
    submission_dates_label, submission_dates_joiner = label_and_joiner("submission_dates_cn", "", "")
    corresponding_author_label, corresponding_author_joiner = label_and_joiner(
        "corresponding_author",
        "通信作者：",
        "",
    )
    funding_cn_label, funding_cn_joiner = label_and_joiner("funding_cn", "基金项目：", "")
    abstract_en_label, abstract_en_joiner = label_and_joiner("abstract_en", "Abstract:", " ")
    keywords_en_label, keywords_en_joiner = label_and_joiner("keywords_en", "Keywords:", " ")
    funding_en_label, funding_en_joiner = label_and_joiner("funding_en", "Foundation Items:", " ")

    return (
        FrontMatterField(
            spec_key="title_cn",
            style_key="title_cn",
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
        ),
        FrontMatterField(
            spec_key="authors_cn",
            style_key="authors_cn",
            mode="join_list",
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            joiner=authors_cn_joiner,
        ),
        FrontMatterField(
            spec_key="affiliations_cn",
            style_key="affiliations_cn",
            mode="list",
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
        ),
        FrontMatterField(
            spec_key="abstract_cn",
            style_key="abstract_cn",
            label=abstract_cn_label,
            label_joiner=abstract_cn_joiner,
        ),
        FrontMatterField(
            spec_key="keywords_cn",
            style_key="keywords_cn",
            mode="keywords",
            label=keywords_cn_label,
            label_joiner=keywords_cn_joiner,
            separator="；",
        ),
        FrontMatterField(
            spec_key="classification_cn",
            style_key="classification_cn",
            label=classification_label,
            label_joiner=classification_joiner,
        ),
        FrontMatterField(
            spec_key="submission_dates_cn",
            style_key="submission_dates_cn",
            label=submission_dates_label or None,
            label_joiner=submission_dates_joiner,
        ),
        FrontMatterField(
            spec_key="corresponding_author",
            style_key="corresponding_author",
            label=corresponding_author_label,
            label_joiner=corresponding_author_joiner,
        ),
        FrontMatterField(
            spec_key="funding_cn",
            style_key="funding_cn",
            label=funding_cn_label,
            label_joiner=funding_cn_joiner,
        ),
        FrontMatterField(
            spec_key="citation_cn",
            style_key="citation_cn",
        ),
        FrontMatterField(
            spec_key="title_en",
            style_key="title_en",
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
        ),
        FrontMatterField(
            spec_key="authors_en",
            style_key="authors_en",
            mode="join_list",
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            joiner=authors_en_joiner,
        ),
        FrontMatterField(
            spec_key="affiliations_en",
            style_key="affiliations_en",
            mode="list",
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
        ),
        FrontMatterField(
            spec_key="abstract_en",
            style_key="abstract_en",
            label=abstract_en_label,
            label_joiner=abstract_en_joiner,
        ),
        FrontMatterField(
            spec_key="keywords_en",
            style_key="keywords_en",
            mode="keywords",
            label=keywords_en_label,
            label_joiner=keywords_en_joiner,
            separator="; ",
        ),
        FrontMatterField(
            spec_key="funding_en",
            style_key="funding_en",
            label=funding_en_label,
            label_joiner=funding_en_joiner,
        ),
        FrontMatterField(
            spec_key="citation_en",
            style_key="citation_en",
        ),
    )


def _build_style_candidates(hints: TemplateHints) -> dict[str, list[str]]:
    style_candidates = {key: list(values) for key, values in BASE_STYLE_CANDIDATES.items()}
    for key, preferred_style in hints.preferred_styles.items():
        _prefer_style_candidate(style_candidates, key, preferred_style)

    cn_front_matter_style = hints.preferred_styles.get("abstract_cn")
    if cn_front_matter_style:
        for key in ("submission_dates_cn", "corresponding_author", "funding_cn"):
            if key not in hints.preferred_styles:
                _prefer_style_candidate(style_candidates, key, cn_front_matter_style)

    en_front_matter_style = hints.preferred_styles.get("abstract_en") or cn_front_matter_style
    if en_front_matter_style and "funding_en" not in hints.preferred_styles:
        _prefer_style_candidate(style_candidates, "funding_en", en_front_matter_style)

    return style_candidates


def _prefer_style_candidate(style_candidates: dict[str, list[str]], key: str, preferred_style: str | None) -> None:
    if not preferred_style or key not in style_candidates:
        return
    style_candidates[key] = [preferred_style] + [
        candidate for candidate in style_candidates[key] if candidate != preferred_style
    ]


def _collect_template_hints(doc: Document) -> TemplateHints:
    paragraphs = [paragraph for paragraph in doc.paragraphs if paragraph.text.strip()]
    preferred_styles: dict[str, str] = {}
    labels: dict[str, str] = {}
    label_joiners: dict[str, str] = {}
    list_joiners: dict[str, str] = {}

    if paragraphs:
        preferred_styles["title_cn"] = paragraphs[0].style.name

    front_matter_boundary = _front_matter_boundary(paragraphs)
    front_matter_paragraphs = paragraphs[:front_matter_boundary]

    _capture_front_matter_field(
        front_matter_paragraphs,
        preferred_styles,
        labels,
        label_joiners,
        spec_key="abstract_cn",
        predicate=lambda text: text.startswith("摘") and "要" in text[:8],
    )
    _capture_front_matter_field(
        front_matter_paragraphs,
        preferred_styles,
        labels,
        label_joiners,
        spec_key="keywords_cn",
        predicate=lambda text: text.startswith("关键词"),
    )
    _capture_front_matter_field(
        front_matter_paragraphs,
        preferred_styles,
        labels,
        label_joiners,
        spec_key="classification_cn",
        predicate=lambda text: text.startswith("中图分类号"),
    )
    _capture_front_matter_field(
        front_matter_paragraphs,
        preferred_styles,
        labels,
        label_joiners,
        spec_key="submission_dates_cn",
        predicate=lambda text: text.startswith("收稿日期"),
    )
    _capture_front_matter_field(
        front_matter_paragraphs,
        preferred_styles,
        labels,
        label_joiners,
        spec_key="corresponding_author",
        predicate=lambda text: text.startswith("通信作者"),
    )
    _capture_front_matter_field(
        front_matter_paragraphs,
        preferred_styles,
        labels,
        label_joiners,
        spec_key="funding_cn",
        predicate=lambda text: text.startswith("基金项目"),
    )
    _capture_front_matter_field(
        front_matter_paragraphs,
        preferred_styles,
        labels,
        label_joiners,
        spec_key="abstract_en",
        predicate=lambda text: text.startswith("Abstract"),
    )
    _capture_front_matter_field(
        front_matter_paragraphs,
        preferred_styles,
        labels,
        label_joiners,
        spec_key="keywords_en",
        predicate=lambda text: text.startswith("Keywords") or text.startswith("Key words"),
    )
    _capture_front_matter_field(
        front_matter_paragraphs,
        preferred_styles,
        labels,
        label_joiners,
        spec_key="funding_en",
        predicate=lambda text: text.startswith("Foundation Items"),
    )

    _capture_named_style(front_matter_paragraphs, preferred_styles, "authors_cn", {"人名", "作者", "CAuthor"})
    _capture_named_style(front_matter_paragraphs, preferred_styles, "affiliations_cn", {"地名", "单位", "CAffiliation"})
    _capture_named_style(front_matter_paragraphs, preferred_styles, "authors_en", {"人名", "Author", "Authors"})
    _capture_named_style(front_matter_paragraphs, preferred_styles, "affiliations_en", {"地名", "Affiliation"})

    if "authors_cn" not in preferred_styles and len(front_matter_paragraphs) > 1:
        preferred_styles["authors_cn"] = front_matter_paragraphs[1].style.name
    if "affiliations_cn" not in preferred_styles and len(front_matter_paragraphs) > 2:
        preferred_styles["affiliations_cn"] = front_matter_paragraphs[2].style.name

    english_title_paragraph = _find_first(
        front_matter_paragraphs,
        lambda paragraph: _looks_like_english_title(paragraph.text.strip()),
    )
    if english_title_paragraph is not None:
        preferred_styles["title_en"] = english_title_paragraph.style.name
        english_index = front_matter_paragraphs.index(english_title_paragraph)
        if "authors_en" not in preferred_styles and english_index + 1 < len(front_matter_paragraphs):
            preferred_styles["authors_en"] = front_matter_paragraphs[english_index + 1].style.name
        if "affiliations_en" not in preferred_styles and english_index + 2 < len(front_matter_paragraphs):
            preferred_styles["affiliations_en"] = front_matter_paragraphs[english_index + 2].style.name

    _capture_list_joiner(front_matter_paragraphs, preferred_styles, list_joiners, "authors_cn")
    _capture_list_joiner(front_matter_paragraphs, preferred_styles, list_joiners, "authors_en")

    heading_1_paragraph = _find_first(paragraphs, lambda paragraph: TOP_LEVEL_HEADING_RE.match(paragraph.text.strip()))
    heading_2_paragraph = _find_first(paragraphs, lambda paragraph: SECOND_LEVEL_HEADING_RE.match(paragraph.text.strip()))
    heading_3_paragraph = _find_first(paragraphs, lambda paragraph: THIRD_LEVEL_HEADING_RE.match(paragraph.text.strip()))
    if heading_1_paragraph is not None:
        preferred_styles["heading_1"] = heading_1_paragraph.style.name
    if heading_2_paragraph is not None:
        preferred_styles["heading_2"] = heading_2_paragraph.style.name
    if heading_3_paragraph is not None:
        preferred_styles["heading_3"] = heading_3_paragraph.style.name

    if heading_1_paragraph is not None:
        start_index = paragraphs.index(heading_1_paragraph) + 1
        body_paragraph = _find_first(
            paragraphs[start_index:],
            lambda paragraph: not TOP_LEVEL_HEADING_RE.match(paragraph.text.strip())
            and not SECOND_LEVEL_HEADING_RE.match(paragraph.text.strip())
            and not THIRD_LEVEL_HEADING_RE.match(paragraph.text.strip())
            and "参考文献" not in paragraph.text.strip(),
        )
        if body_paragraph is not None:
            preferred_styles["body"] = body_paragraph.style.name

    reference_heading_paragraph = _find_first(
        paragraphs,
        lambda paragraph: _is_reference_heading(paragraph.text.strip()),
    )
    reference_heading_text = None
    reference_numbering_mode = "style"
    if reference_heading_paragraph is not None:
        preferred_styles["reference_heading"] = reference_heading_paragraph.style.name
        reference_heading_text = reference_heading_paragraph.text.strip()
        reference_index = paragraphs.index(reference_heading_paragraph)
        reference_item_paragraph = _find_first(
            paragraphs[reference_index + 1 :],
            lambda paragraph: bool(paragraph.text.strip()) and "作者简介" not in paragraph.text.strip(),
        )
        if reference_item_paragraph is not None:
            preferred_styles["reference_cn_text"] = reference_item_paragraph.style.name
            preferred_styles["reference_en_text"] = reference_item_paragraph.style.name
            if REFERENCE_ITEM_RE.match(reference_item_paragraph.text.strip()) and "<w:numPr" not in reference_item_paragraph._p.xml:
                reference_numbering_mode = "manual_brackets"

    acknowledgement_paragraph = _find_first(
        paragraphs,
        lambda paragraph: _is_acknowledgement_heading(paragraph.text.strip()),
    )
    acknowledgement_heading_text = None
    if acknowledgement_paragraph is not None:
        preferred_styles["ack_heading"] = acknowledgement_paragraph.style.name
        acknowledgement_heading_text = acknowledgement_paragraph.text.strip()

    return TemplateHints(
        preferred_styles=preferred_styles,
        labels=labels,
        label_joiners=label_joiners,
        list_joiners=list_joiners,
        inferred_body_columns=_detect_body_columns_from_doc(doc),
        reference_heading_text=reference_heading_text,
        acknowledgement_heading_text=acknowledgement_heading_text,
        reference_numbering_mode=reference_numbering_mode,
        header_sample_text=_first_non_empty_header_text(doc),
    )


def _capture_front_matter_field(
    paragraphs: list,
    preferred_styles: dict[str, str],
    labels: dict[str, str],
    label_joiners: dict[str, str],
    *,
    spec_key: str,
    predicate,
) -> None:
    paragraph = _find_first(paragraphs, lambda candidate: predicate(candidate.text.strip()))
    if paragraph is None:
        return
    preferred_styles[spec_key] = paragraph.style.name
    label, joiner = _extract_label_hint(paragraph.text.strip())
    if label is not None:
        labels[spec_key] = label
        label_joiners[spec_key] = joiner


def _capture_named_style(paragraphs: list, preferred_styles: dict[str, str], spec_key: str, names: set[str]) -> None:
    paragraph = _find_first(paragraphs, lambda candidate: candidate.style.name in names)
    if paragraph is not None:
        preferred_styles[spec_key] = paragraph.style.name


def _capture_list_joiner(
    paragraphs: list,
    preferred_styles: dict[str, str],
    list_joiners: dict[str, str],
    spec_key: str,
) -> None:
    style_name = preferred_styles.get(spec_key)
    if not style_name:
        return
    paragraph = _find_first(paragraphs, lambda candidate: candidate.style.name == style_name and candidate.text.strip())
    if paragraph is None:
        return
    joiner = _infer_author_joiner(paragraph.text.strip())
    if joiner:
        list_joiners[spec_key] = joiner


def _infer_author_joiner(text: str) -> str | None:
    for joiner in ("，", ", ", ",", "；", "; "):
        pieces = [part.strip() for part in text.split(joiner) if part.strip()]
        if len(pieces) >= 2:
            return joiner
    return None


def _extract_label_hint(text: str) -> tuple[str | None, str]:
    for separator in ("：", ":"):
        index = text.find(separator)
        if 0 < index <= 20:
            label = text[: index + 1]
            joiner = " " if len(text) > index + 1 and text[index + 1] == " " else ""
            return label, joiner
    return None, ""


def _front_matter_boundary(paragraphs: list) -> int:
    heading_index = next(
        (
            index
            for index, paragraph in enumerate(paragraphs)
            if TOP_LEVEL_HEADING_RE.match(paragraph.text.strip()) or SECOND_LEVEL_HEADING_RE.match(paragraph.text.strip())
        ),
        None,
    )
    if heading_index is None:
        return min(len(paragraphs), 12)
    return heading_index


def _looks_like_english_title(text: str) -> bool:
    stripped = text.strip()
    if len(stripped) < 12:
        return False
    if stripped.startswith("Abstract") or stripped.startswith("Keywords") or stripped.startswith("Key words"):
        return False
    ascii_letters = sum(1 for char in stripped if char.isascii() and char.isalpha())
    return ascii_letters >= max(8, len(stripped) // 3)


def _is_reference_heading(text: str) -> bool:
    stripped = text.strip()
    lower = stripped.lower().rstrip(":：")
    return "参考文献" in stripped or lower in {"references", "reference", "bibliography"}


def _is_acknowledgement_heading(text: str) -> bool:
    stripped = text.strip()
    lower = stripped.lower().rstrip(":：")
    return "致谢" in stripped or lower in {"acknowledgements", "acknowledgments"}


def _find_first(items: list, predicate):
    for item in items:
        if predicate(item):
            return item
    return None


def _detect_primary_language(spec: dict[str, Any]) -> str:
    chinese_keys = ("title_cn", "authors_cn", "affiliations_cn", "abstract_cn", "keywords_cn")
    english_keys = ("title_en", "authors_en", "affiliations_en", "abstract_en", "keywords_en")
    chinese_score = sum(1 for key in chinese_keys if spec.get(key))
    english_score = sum(1 for key in english_keys if spec.get(key))
    if english_score > chinese_score:
        return "en"
    return "cn"


def _detect_body_columns_from_doc(doc: Document) -> int | None:
    if len(doc.sections) < 2:
        return None

    primary_columns = _section_columns(doc.sections[0]._sectPr)
    counts: dict[int, int] = {}
    for section in doc.sections[1:]:
        columns = _section_columns(section._sectPr)
        if columns != primary_columns:
            counts[columns] = counts.get(columns, 0) + 1

    if not counts:
        return None

    return sorted(counts.items(), key=lambda item: (-item[1], item[0]))[0][0]


def _section_columns(sect_pr) -> int:
    cols = sect_pr.find(qn("w:cols"))
    if cols is None:
        return 1
    try:
        return int(cols.get(qn("w:num")) or "1")
    except (TypeError, ValueError):
        return 1


def _first_non_empty_header_text(doc: Document) -> str | None:
    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            text = paragraph.text.strip()
            if text:
                return text
    return None


def _update_running_headers(template_doc: Document, spec: dict[str, Any], output_path: str | Path) -> bool:
    if _first_non_empty_header_text(template_doc) is None:
        return False

    header_text = _build_running_header_text(spec)
    if not header_text:
        return False

    output_doc = Document(str(output_path))
    changed = False
    for index, section in enumerate(output_doc.sections):
        target = _find_first(section.header.paragraphs, lambda paragraph: paragraph.text.strip())
        if target is None:
            target = section.header.paragraphs[0] if section.header.paragraphs else section.header.add_paragraph()
        template_section = template_doc.sections[min(index, len(template_doc.sections) - 1)]
        template_target = _find_first(template_section.header.paragraphs, lambda paragraph: paragraph.text.strip())
        skeleton = template_target or target
        wrapped_header_text = _wrap_header_text_like_paragraph(skeleton.text, header_text)

        if template_target is not None:
            target.style = template_target.style.name
            target.alignment = template_target.alignment

        if target.text != wrapped_header_text:
            target.text = wrapped_header_text
            changed = True

    if changed:
        output_doc.save(str(output_path))
    return changed


def _wrap_header_text_like_paragraph(template_text: str, header_text: str) -> str:
    leading_tabs = len(template_text) - len(template_text.lstrip("\t"))
    trailing_tabs = len(template_text) - len(template_text.rstrip("\t"))
    return f"{'\t' * leading_tabs}{header_text}{'\t' * trailing_tabs}"


def _build_running_header_text(spec: dict[str, Any]) -> str:
    title_cn = str(spec.get("title_cn") or "").strip()
    title_en = str(spec.get("title_en") or "").strip()
    authors_cn = _normalize_list(spec.get("authors_cn"))
    authors_en = _normalize_list(spec.get("authors_en"))

    if title_cn:
        if authors_cn:
            author = _short_author_name(authors_cn[0])
            return f"{author}{'等' if len(authors_cn) > 1 else ''}：{title_cn}"
        return title_cn

    if title_en:
        if authors_en:
            author = _short_author_name(authors_en[0])
            return f"{author}{' et al.' if len(authors_en) > 1 else ''}: {title_en}"
        return title_en

    return ""


def _normalize_list(value: Any) -> list[str]:
    if value is None:
        return []
    if isinstance(value, str):
        text = value.strip()
        return [text] if text else []
    return [str(item).strip() for item in value if str(item).strip()]


def _short_author_name(name: str) -> str:
    return re.sub(r"[\d*†‡]+$", "", name.strip())


def _serialize_hints(hints: TemplateHints) -> dict[str, Any]:
    return {
        "preferred_styles": hints.preferred_styles,
        "labels": hints.labels,
        "label_joiners": hints.label_joiners,
        "list_joiners": hints.list_joiners,
        "inferred_body_columns": hints.inferred_body_columns,
        "reference_heading_text": hints.reference_heading_text,
        "acknowledgement_heading_text": hints.acknowledgement_heading_text,
        "reference_numbering_mode": hints.reference_numbering_mode,
        "header_sample_text": hints.header_sample_text,
    }


def _prepare_spec(spec: dict[str, Any], hints: TemplateHints) -> dict[str, Any]:
    prepared = dict(spec)
    prepared["authors_cn"] = _normalize_listish_value(
        spec.get("authors_cn"),
        separators=(r"《\s*[,，;；]\s*》", r"[；;，、\n]+"),
    )
    prepared["authors_en"] = _normalize_listish_value(
        spec.get("authors_en"),
        separators=(r"《\s*[,，;；]\s*》", r"[；;\n]+"),
    )
    prepared["keywords_cn"] = _normalize_listish_value(
        spec.get("keywords_cn"),
        separators=(r"《\s*[,，;；]\s*》", r"[；;，,、\n]+"),
    )
    prepared["keywords_en"] = _normalize_listish_value(
        spec.get("keywords_en"),
        separators=(r"《\s*[,，;；]\s*》", r"[；;,，\n]+"),
    )
    submission_dates_cn = _compose_submission_dates_cn(spec, hints)
    if submission_dates_cn:
        prepared["submission_dates_cn"] = submission_dates_cn
    return prepared


def _normalize_listish_value(value: Any, *, separators: tuple[str, ...]) -> list[str]:
    if value is None:
        return []

    raw_items = value if isinstance(value, list) else [value]
    pattern = re.compile("|".join(f"(?:{separator})" for separator in separators))
    normalized: list[str] = []

    for raw_item in raw_items:
        text = str(raw_item).strip()
        if not text:
            continue
        parts = [part.strip() for part in pattern.split(text) if part.strip()]
        if parts:
            normalized.extend(parts)
        else:
            normalized.append(text)
    return normalized


def _compose_submission_dates_cn(spec: dict[str, Any], hints: TemplateHints) -> str:
    explicit = str(spec.get("submission_dates_cn") or "").strip()
    if explicit:
        return explicit

    received_date = str(spec.get("received_date") or "").strip()
    revised_date = str(spec.get("revised_date") or "").strip()
    if not received_date and not revised_date:
        return ""

    label = hints.labels.get("submission_dates_cn", "")
    if label:
        received_prefix = label
    else:
        received_prefix = "收稿日期："
    revised_prefix = "修回日期："

    parts: list[str] = []
    if received_date:
        parts.append(f"{received_prefix}{received_date}")
    if revised_date:
        parts.append(f"{revised_prefix}{revised_date}")
    return "；".join(parts)


def main() -> None:
    parser = argparse.ArgumentParser(description="Inspect and write generic DOCX templates.")
    subparsers = parser.add_subparsers(dest="command", required=True)

    inspect_parser = subparsers.add_parser("inspect", help="Inspect template styles and structure.")
    inspect_parser.add_argument("--template", required=True, help="Path to the .docx template.")
    inspect_parser.add_argument("--out", help="Optional JSON output path.")

    generate_parser = subparsers.add_parser("generate", help="Generate a .docx from a JSON spec.")
    generate_parser.add_argument("--template", required=True, help="Path to the .docx template.")
    generate_parser.add_argument("--spec", required=True, help="Path to the JSON spec.")
    generate_parser.add_argument("--output", required=True, help="Path to the output .docx.")

    args = parser.parse_args()

    if args.command == "inspect":
        result = inspect_template(args.template)
        text = json.dumps(result, ensure_ascii=False, indent=2)
        if args.out:
            Path(args.out).write_text(text, encoding="utf-8")
        print(text)
        return

    result = generate_document(
        template_path=args.template,
        spec=load_json(args.spec),
        output_path=args.output,
    )
    print(json.dumps(result, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
