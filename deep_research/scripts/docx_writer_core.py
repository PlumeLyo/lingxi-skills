from __future__ import annotations

import json
import re
from copy import deepcopy
from dataclasses import dataclass
from pathlib import Path
from typing import Any

from docx import Document
from docx.document import Document as DocxDocument
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Mm, Pt


MANUAL_HEADING_NUMBER_RE = re.compile(
    r"^\s*[1-9]\d?(?:\.\d+){0,2}(?:[、.)．])?\s+(.+?)\s*$"
)
BOLD_SPAN_RE = re.compile(r"(\*\*.+?\*\*)")
BULLET_ITEM_RE = re.compile(r"^[-*+]\s+(.*\S)\s*$")
ORDERED_ITEM_RE = re.compile(r"^(\d+[.)])\s+(.*\S)\s*$")


@dataclass(frozen=True)
class FrontMatterField:
    spec_key: str
    style_key: str
    mode: str = "text"
    alignment: Any = None
    label: str | None = None
    label_joiner: str = "\t"
    separator: str = ""
    joiner: str = "  "


@dataclass(frozen=True)
class WriterConfig:
    style_candidates: dict[str, list[str]]
    front_matter_fields: tuple[FrontMatterField, ...]
    body_columns: int | None = None
    acknowledgement_heading: str = "致  谢"
    reference_heading: str = "参 考 文 献"
    appendix_default_title: str = "附录"
    reference_numbering_mode: str = "style"


def inspect_template(
    template_path: str | Path,
    style_candidates: dict[str, list[str]],
    sample_limit: int = 40,
) -> dict[str, Any]:
    doc = Document(str(template_path))
    samples: list[dict[str, Any]] = []
    for index, paragraph in enumerate(doc.paragraphs, start=1):
        text = paragraph.text.strip()
        if not text:
            continue
        samples.append(
            {
                "index": index,
                "style": paragraph.style.name,
                "text": text[:160],
            }
        )
        if len(samples) >= sample_limit:
            break

    return {
        "template_path": str(template_path),
        "paragraph_count": len(doc.paragraphs),
        "table_count": len(doc.tables),
        "section_count": len(doc.sections),
        "style_map": build_style_map(doc, style_candidates),
        "sample_paragraphs": samples,
    }


def generate_document(
    template_path: str | Path,
    spec: dict[str, Any],
    output_path: str | Path,
    config: WriterConfig,
) -> dict[str, Any]:
    doc = Document(str(template_path))
    style_map = build_style_map(doc, config.style_candidates)
    primary_sect_pr = _primary_sect_pr(doc)
    body_sect_pr = None
    if config.body_columns is not None:
        body_sect_pr = _find_sect_pr_by_columns(doc, columns=config.body_columns)
    _clear_body(doc, sect_pr_template=primary_sect_pr)

    doc.core_properties.title = str(spec.get("title_cn") or spec.get("title_en") or "")
    doc.core_properties.author = "; ".join(_normalize_list(spec.get("authors_cn")))

    _write_front_matter(doc, spec, style_map, config)

    has_main_content = bool(
        _iter_sections(spec)
        or str(spec.get("acknowledgements", "")).strip()
        or spec.get("references")
        or spec.get("appendices")
    )
    if (
        body_sect_pr is not None
        and has_main_content
        and _sectpr_columns(body_sect_pr) != _sectpr_columns(primary_sect_pr)
    ):
        _switch_to_layout(doc, primary_sect_pr, body_sect_pr)

    _write_sections(doc, spec, style_map)
    _write_acknowledgements(doc, spec, style_map, config)
    _write_references(doc, spec, style_map, config)
    _write_appendices(doc, spec, style_map, config)

    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output))
    return {
        "output_path": str(output),
        "style_map": style_map,
        "section_count": len(_iter_sections(spec)),
        "reference_count": len(_normalize_references(spec.get("references"))),
    }


def load_json(path: str | Path) -> dict[str, Any]:
    return json.loads(Path(path).read_text(encoding="utf-8"))


def build_style_map(doc: DocxDocument, style_candidates: dict[str, list[str]]) -> dict[str, str]:
    return {key: _pick_style(doc, key, style_candidates) for key in style_candidates}


def _style_names(doc: DocxDocument) -> set[str]:
    return {style.name for style in doc.styles if style.type == WD_STYLE_TYPE.PARAGRAPH}


def _pick_style(doc: DocxDocument, key: str, style_candidates: dict[str, list[str]]) -> str:
    available = _style_names(doc)
    for candidate in style_candidates[key]:
        if candidate in available:
            return candidate
    return "Normal"


def _primary_sect_pr(doc: DocxDocument):
    # Preserve the template's primary page settings and header/footer linkage.
    return deepcopy(doc.sections[0]._sectPr)


def _sectpr_columns(sect_pr) -> int:
    cols = sect_pr.find(qn("w:cols"))
    if cols is None:
        return 1
    try:
        return int(cols.get(qn("w:num")) or "1")
    except (TypeError, ValueError):
        return 1


def _find_sect_pr_by_columns(doc: DocxDocument, columns: int, fallback_index: int = 0):
    for section in doc.sections:
        if _sectpr_columns(section._sectPr) == columns:
            return deepcopy(section._sectPr)
    return deepcopy(doc.sections[fallback_index]._sectPr)


def _clear_body(doc: DocxDocument, sect_pr_template=None) -> None:
    body = doc._element.body
    children = list(body.iterchildren())
    sect_pr = None
    for child in children:
        if child.tag.endswith("sectPr"):
            sect_pr = child
            break
    for child in children:
        if child is not sect_pr:
            body.remove(child)
    if sect_pr_template is not None:
        if sect_pr is not None:
            body.remove(sect_pr)
        body.append(deepcopy(sect_pr_template))


def _set_body_section_layout(doc: DocxDocument, sect_pr_template) -> None:
    body = doc._element.body
    sect_pr = body.sectPr
    parent = sect_pr.getparent()
    parent.replace(sect_pr, deepcopy(sect_pr_template))


def _switch_to_layout(doc: DocxDocument, current_sect_pr_template, next_sect_pr_template) -> None:
    paragraph = doc.add_paragraph()
    p_pr = paragraph._p.get_or_add_pPr()
    sect_pr = OxmlElement("w:sectPr")
    for child in deepcopy(current_sect_pr_template):
        sect_pr.append(child)
    p_pr.append(sect_pr)
    _set_body_section_layout(doc, next_sect_pr_template)


def _style_has_numbering(doc: DocxDocument, style_name: str) -> bool:
    try:
        style = doc.styles[style_name]
    except KeyError:
        return False
    return "<w:numPr" in style.element.xml


def _style_numbering(doc: DocxDocument, style_name: str) -> tuple[str | None, str | None]:
    try:
        style = doc.styles[style_name]
    except KeyError:
        return None, None
    p_pr = style.element.find(qn("w:pPr"))
    if p_pr is None:
        return None, None
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        return None, None
    ilvl = num_pr.find(qn("w:ilvl"))
    num_id = num_pr.find(qn("w:numId"))
    return (
        ilvl.get(qn("w:val")) if ilvl is not None else None,
        num_id.get(qn("w:val")) if num_id is not None else None,
    )


def _style_paragraph_layout(doc: DocxDocument, style_name: str) -> dict[str, Any]:
    try:
        style = doc.styles[style_name]
    except KeyError:
        return {}
    p_pr = style.element.find(qn("w:pPr"))
    if p_pr is None:
        return {}

    layout: dict[str, Any] = {}
    for tag in ("w:tabs", "w:ind"):
        child = p_pr.find(qn(tag))
        if child is not None:
            layout[tag] = deepcopy(child)
    return layout


def _numbering_level_layout(doc: DocxDocument, num_id: str | None, ilvl: str | int = 0) -> dict[str, Any]:
    if num_id is None:
        return {}

    try:
        numbering = doc.part.numbering_part.element
    except Exception:
        return {}

    ilvl = str(ilvl)
    num_xpath = f'.//w:num[@w:numId="{num_id}"]'
    num_matches = numbering.xpath(num_xpath)
    if not num_matches:
        return {}

    abstract_num_id = None
    for child in num_matches[0]:
        if child.tag == qn("w:abstractNumId"):
            abstract_num_id = child.get(qn("w:val"))
            break
    if abstract_num_id is None:
        return {}

    abs_xpath = f'.//w:abstractNum[@w:abstractNumId="{abstract_num_id}"]/w:lvl[@w:ilvl="{ilvl}"]/w:pPr'
    abs_matches = numbering.xpath(abs_xpath)
    if not abs_matches:
        return {}

    layout: dict[str, Any] = {}
    for tag in ("w:tabs", "w:ind"):
        child = abs_matches[0].find(qn(tag))
        if child is not None:
            layout[tag] = deepcopy(child)
    return layout


def _strip_manual_heading_number(text: str) -> str:
    match = MANUAL_HEADING_NUMBER_RE.match(text)
    if match:
        return match.group(1)
    return text.strip()


def _append_inline_runs(paragraph, text: str) -> None:
    for segment in BOLD_SPAN_RE.split(text):
        if not segment:
            continue
        if segment.startswith("**") and segment.endswith("**") and len(segment) >= 4:
            run = paragraph.add_run(segment[2:-2])
            run.bold = True
            continue
        paragraph.add_run(segment)


def _iter_text_fragments(text: str) -> list[dict[str, str]]:
    fragments: list[dict[str, str]] = []
    paragraph_lines: list[str] = []

    def flush_paragraph() -> None:
        if not paragraph_lines:
            return
        fragments.append({"type": "paragraph", "text": " ".join(paragraph_lines).strip()})
        paragraph_lines.clear()

    for raw_line in text.replace("\r\n", "\n").split("\n"):
        line = raw_line.strip()
        if not line:
            flush_paragraph()
            continue

        bullet_match = BULLET_ITEM_RE.match(line)
        if bullet_match:
            flush_paragraph()
            fragments.append({"type": "bullet", "text": bullet_match.group(1).strip()})
            continue

        ordered_match = ORDERED_ITEM_RE.match(line)
        if ordered_match:
            flush_paragraph()
            fragments.append(
                {
                    "type": "ordered",
                    "text": f"{ordered_match.group(1)} {ordered_match.group(2).strip()}",
                }
            )
            continue

        paragraph_lines.append(line)

    flush_paragraph()
    return fragments


def _set_paragraph_numbering(paragraph, num_id: str, ilvl: str | int = 0) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    existing = p_pr.find(qn("w:numPr"))
    if existing is not None:
        p_pr.remove(existing)
    num_pr = OxmlElement("w:numPr")
    ilvl_el = OxmlElement("w:ilvl")
    ilvl_el.set(qn("w:val"), str(ilvl))
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl_el)
    num_pr.append(num_id_el)
    p_pr.append(num_pr)


def _apply_paragraph_layout(paragraph, layout: dict[str, Any]) -> None:
    if not layout:
        return
    p_pr = paragraph._p.get_or_add_pPr()
    for tag, child in layout.items():
        existing = p_pr.find(qn(tag))
        if existing is not None:
            p_pr.remove(existing)
        p_pr.append(deepcopy(child))


def _add_paragraph(
    doc: DocxDocument,
    text: str,
    style_name: str,
    *,
    alignment=None,
    prefix: str = "",
):
    paragraph = doc.add_paragraph(style=style_name)
    if alignment is not None:
        paragraph.alignment = alignment
    if prefix:
        paragraph.add_run(prefix)
    lines = text.split("\n")
    for index, line in enumerate(lines):
        if index:
            paragraph.add_run().add_break()
        _append_inline_runs(paragraph, line)
    return paragraph


def _add_heading_paragraph(doc: DocxDocument, text: str, style_name: str) -> None:
    heading_text = _strip_manual_heading_number(text) if _style_has_numbering(doc, style_name) else text
    _add_paragraph(doc, heading_text, style_name)


def _write_body_text(doc: DocxDocument, text: str, style_name: str) -> None:
    for fragment in _iter_text_fragments(text):
        if fragment["type"] == "bullet":
            _add_paragraph(doc, fragment["text"], style_name, prefix="• ")
            continue
        _add_paragraph(doc, fragment["text"], style_name)


def _normalize_list(value: Any) -> list[str]:
    if value is None:
        return []
    if isinstance(value, str):
        return [value]
    return [str(item) for item in value if str(item).strip()]


def _split_paragraphs(value: Any) -> list[str]:
    if value is None:
        return []
    if isinstance(value, str):
        text = value.replace("\r\n", "\n").strip()
        if not text:
            return []
        return [chunk.strip() for chunk in re.split(r"\n\s*\n", text) if chunk.strip()]
    if isinstance(value, list):
        paragraphs: list[str] = []
        for item in value:
            paragraphs.extend(_split_paragraphs(item))
        return paragraphs
    if isinstance(value, dict):
        for key in ("text", "content", "paragraphs"):
            if key in value:
                return _split_paragraphs(value.get(key))
        return []
    return _split_paragraphs(str(value))


def _join_keywords(keywords: Any, separator: str) -> str:
    if isinstance(keywords, list):
        return separator.join(str(item).strip() for item in keywords if str(item).strip())
    return str(keywords).strip()


def _ensure_label(text: str, label: str, joiner: str = "\t") -> str:
    stripped = text.lstrip()
    if stripped.startswith(label):
        return text
    return f"{label}{joiner}{text}"


def _normalize_references(raw_references: Any) -> list[str]:
    references: list[str] = []
    if raw_references is None:
        return references
    if not isinstance(raw_references, list):
        raw_references = [raw_references]
    for item in raw_references:
        if isinstance(item, dict):
            text = str(item.get("content") or item.get("text") or "").strip()
        else:
            text = str(item).strip()
        if text:
            references.append(text)
    return references


def _strip_reference_number(text: str) -> str:
    return re.sub(r"^\s*\[\d+\]\s*", "", text).strip()


def _write_front_matter(
    doc: DocxDocument,
    spec: dict[str, Any],
    style_map: dict[str, str],
    config: WriterConfig,
) -> None:
    for field in config.front_matter_fields:
        if field.mode == "list":
            for item in _normalize_list(spec.get(field.spec_key)):
                text = _apply_label(item, field)
                _add_paragraph(doc, text, style_map[field.style_key], alignment=field.alignment)
            continue

        if field.mode == "join_list":
            items = _normalize_list(spec.get(field.spec_key))
            if not items:
                continue
            text = field.joiner.join(items)
        elif field.mode == "keywords":
            text = _join_keywords(spec.get(field.spec_key, []), field.separator)
        else:
            text = str(spec.get(field.spec_key, "")).strip()

        text = text.strip()
        if not text:
            continue
        text = _apply_label(text, field)
        _add_paragraph(doc, text, style_map[field.style_key], alignment=field.alignment)


def _apply_label(text: str, field: FrontMatterField) -> str:
    if field.label:
        return _ensure_label(text, field.label, field.label_joiner)
    return text


def _iter_blocks(section: dict[str, Any]) -> list[dict[str, Any]]:
    if section.get("blocks"):
        return list(section["blocks"])
    paragraphs_value = (
        section.get("paragraphs")
        if section.get("paragraphs") is not None
        else section.get("content")
    )
    paragraphs = _split_paragraphs(paragraphs_value)
    return [{"type": "paragraph", "text": paragraph} for paragraph in paragraphs]


def _write_table(doc: DocxDocument, rows: list[list[Any]]) -> None:
    if not rows:
        return
    max_cols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=max_cols)
    if "Table Grid" in {style.name for style in doc.styles}:
        table.style = "Table Grid"
    for row_index, row in enumerate(rows):
        for col_index in range(max_cols):
            cell_text = str(row[col_index]) if col_index < len(row) else ""
            table.cell(row_index, col_index).text = cell_text


def _write_image(doc: DocxDocument, block: dict[str, Any], style_map: dict[str, str]) -> None:
    raw_path = str(block.get("path") or block.get("src") or "").strip()
    if not raw_path:
        return

    image_path = Path(raw_path).expanduser()
    if not image_path.exists():
        raise FileNotFoundError(f"Image block path does not exist: {image_path}")

    alignment = _parse_alignment(block.get("alignment"))
    paragraph = doc.add_paragraph(style=style_map["body"])
    paragraph.alignment = alignment

    width = _resolve_length(block, "width")
    height = _resolve_length(block, "height")
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=width, height=height)

    caption = str(block.get("caption", "")).strip()
    if caption:
        caption_paragraph = doc.add_paragraph(style=style_map["body"])
        caption_paragraph.alignment = alignment
        _append_inline_runs(caption_paragraph, caption)


def _parse_alignment(value: Any):
    normalized = str(value or "center").strip().lower()
    if normalized in {"left", "start"}:
        return WD_ALIGN_PARAGRAPH.LEFT
    if normalized in {"right", "end"}:
        return WD_ALIGN_PARAGRAPH.RIGHT
    if normalized in {"justify", "justified"}:
        return WD_ALIGN_PARAGRAPH.JUSTIFY
    return WD_ALIGN_PARAGRAPH.CENTER


def _resolve_length(block: dict[str, Any], prefix: str):
    if block.get(f"{prefix}_inches") is not None:
        return Inches(float(block[f"{prefix}_inches"]))
    if block.get(f"{prefix}_cm") is not None:
        return Cm(float(block[f"{prefix}_cm"]))
    if block.get(f"{prefix}_mm") is not None:
        return Mm(float(block[f"{prefix}_mm"]))
    if block.get(f"{prefix}_pt") is not None:
        return Pt(float(block[f"{prefix}_pt"]))
    return None


def _section_level(section: dict[str, Any], default_level: int = 1) -> int:
    raw_level = section.get("level", section.get("heading_level", default_level))
    try:
        level = int(raw_level)
    except (TypeError, ValueError):
        level = default_level
    return max(1, min(level, 3))


def _section_title(section: dict[str, Any], level: int) -> str:
    candidate_keys = ["title", f"heading_{level}", "heading"]
    for key in candidate_keys:
        value = str(section.get(key, "")).strip()
        if value:
            return value
    return ""


def _section_children(section: dict[str, Any]) -> list[dict[str, Any]]:
    children = section.get("subsections")
    if children is None:
        children = section.get("sub_sections")
    return list(children or [])


def _iter_sections(spec: dict[str, Any]) -> list[dict[str, Any]]:
    sections = spec.get("sections")
    if isinstance(sections, list):
        return sections
    body = spec.get("body")
    if isinstance(body, list):
        return body
    return []


def _write_section(doc: DocxDocument, section: dict[str, Any], style_map: dict[str, str], default_level: int = 1) -> None:
    level = _section_level(section, default_level)
    title = _section_title(section, level)
    if title:
        _add_heading_paragraph(doc, title, style_map[f"heading_{level}"])

    for block in _iter_blocks(section):
        block_type = block.get("type", "paragraph")
        if block_type == "paragraph":
            text = str(block.get("text", "")).strip()
            if text:
                _write_body_text(doc, text, style_map["body"])
        elif block_type == "subheading":
            text = str(block.get("text", "")).strip()
            if text:
                _add_heading_paragraph(doc, text, style_map[f"heading_{min(level + 1, 3)}"])
        elif block_type == "table":
            rows = block.get("rows", [])
            _write_table(doc, rows)
        elif block_type == "image":
            _write_image(doc, block, style_map)

    for child in _section_children(section):
        _write_section(doc, child, style_map, default_level=min(level + 1, 3))


def _write_sections(doc: DocxDocument, spec: dict[str, Any], style_map: dict[str, str]) -> None:
    for section in _iter_sections(spec):
        _write_section(doc, section, style_map)


def _write_acknowledgements(
    doc: DocxDocument,
    spec: dict[str, Any],
    style_map: dict[str, str],
    config: WriterConfig,
) -> None:
    text = str(spec.get("acknowledgements", "")).strip()
    if not text:
        return
    _add_paragraph(doc, config.acknowledgement_heading, style_map["ack_heading"])
    _write_body_text(doc, text, style_map["body"])


def _write_references(
    doc: DocxDocument,
    spec: dict[str, Any],
    style_map: dict[str, str],
    config: WriterConfig,
) -> None:
    references = _normalize_references(spec.get("references"))
    if not references:
        return
    _add_paragraph(doc, config.reference_heading, style_map["reference_heading"])
    if config.reference_numbering_mode == "manual_brackets":
        numbering_style = _pick_manual_reference_style(doc, style_map)
        reference_layout = _style_paragraph_layout(doc, numbering_style)
        for index, reference in enumerate(references, start=1):
            paragraph = _add_paragraph(
                doc,
                f"[{index}]\t{_strip_reference_number(reference)}",
                numbering_style,
            )
            _apply_paragraph_layout(paragraph, reference_layout)
        return

    numbering_style = style_map["reference_en_text"]
    if not _style_has_numbering(doc, numbering_style):
        numbering_style = style_map["reference_cn_text"]

    reference_ilvl, reference_num_id = _style_numbering(doc, numbering_style)
    if reference_num_id is None:
        reference_ilvl, reference_num_id = _style_numbering(doc, style_map["reference_cn_text"])
    reference_layout = _numbering_level_layout(doc, reference_num_id, reference_ilvl or 0)
    if not reference_layout:
        reference_layout = _style_paragraph_layout(doc, numbering_style)
    for reference in references:
        paragraph = _add_paragraph(doc, _strip_reference_number(reference), numbering_style)
        if reference_num_id is not None:
            _set_paragraph_numbering(paragraph, reference_num_id, reference_ilvl or 0)
        _apply_paragraph_layout(paragraph, reference_layout)


def _pick_manual_reference_style(doc: DocxDocument, style_map: dict[str, str]) -> str:
    candidates = [style_map["reference_cn_text"], style_map["reference_en_text"]]
    for style_name in candidates:
        if style_name != "Normal" and not _style_has_numbering(doc, style_name):
            return style_name
    for style_name in candidates:
        if style_name != "Normal":
            return style_name
    return "Normal"


def _write_appendices(
    doc: DocxDocument,
    spec: dict[str, Any],
    style_map: dict[str, str],
    config: WriterConfig,
) -> None:
    for appendix in spec.get("appendices", []):
        title = str(appendix.get("title", "")).strip() or config.appendix_default_title
        _add_heading_paragraph(doc, title, style_map["appendix_heading"])
        for paragraph in _normalize_list(appendix.get("paragraphs")):
            _write_body_text(doc, paragraph, style_map["body"])
